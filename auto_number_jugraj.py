import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

folder_path = r"C:\Users\user\OneDrive\Desktop\Jugraj"

# Patterns for identifying question lines and options
question_pattern = re.compile(r"^\s*(?:\d+\.\s*|Q[\.^#\d]*\.?\s*|\^[#\d]+\.\s*)", re.IGNORECASE)
option_pattern = re.compile(r"^\s*(?:\([A-D]\)|[A-D]\.|[1-4]\)|\([1-4]\))", re.IGNORECASE)

def remove_prefix_across_runs(para, prefix_len):
    """Remove prefix_len characters (including old numbers) from the start across runs."""
    to_remove = prefix_len
    removed_text = ""
    for run in para.runs:
        if to_remove <= 0:
            break
        text = run.text or ""
        if not text:
            continue
        if len(text) <= to_remove:
            removed_text += text
            run.text = ""
            to_remove -= len(text)
        else:
            removed_text += text[:to_remove]
            run.text = text[to_remove:]
            to_remove = 0
            break
    return removed_text  # return what was removed (so we can see if it had tabs/spaces)

for filename in os.listdir(folder_path):
    if not filename.lower().endswith(".docx"):
        continue

    file_path = os.path.join(folder_path, filename)
    print(f"Processing: {file_path}")
    doc = Document(file_path)
    question_count = 1
    changed = False

    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs).lstrip()
        if not full_text:
            continue
        if option_pattern.match(full_text):
            continue

        match = question_pattern.match(full_text)
        if not match:
            continue

        old_prefix = match.group(0)
        prefix_len = len(old_prefix)
        new_prefix = f"{question_count}. "

        # Remove old prefix but store what it contained (for tab recovery)
        removed_text = remove_prefix_across_runs(para, prefix_len)

        # Count how many tabs/spaces were originally after prefix
        tab_count = removed_text.count("\t")
        space_count = len(re.findall(r" +", removed_text))  # total spaces
        # Add back roughly the same spacing (tabs have priority)
        gap = "\t" * tab_count if tab_count else (" " * min(space_count, 4) or "\t")

        # Prepare final prefix + preserved spacing
        new_prefix_full = new_prefix + gap

        if para.runs:
            first_run = para.runs[0]
            # Prepend new prefix to text
            first_run.text = new_prefix_full + first_run.text
            # Style prefix run
            first_run.bold = True
            first_run.font.name = 'Times New Roman'
            first_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            first_run.font.size = Pt(11)
        else:
            run = para.add_run(new_prefix_full)
            run.bold = True
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(11)

        question_count += 1
        changed = True

    if changed:
        doc.save(file_path)
        print(f"✅ Updated → {filename}")
    else:
        print(f"⚠️ No changes needed: {filename}")

print("\n🎯 All files processed — old numbers replaced, formatting & tabs preserved!")
